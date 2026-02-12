# -*- coding: utf-8 -*-
"""
SRD Abstracts – Reviewer Assignment & DOCX Processor
Optimized for Stability and Low RAM Usage.
"""

import streamlit as st
import pandas as pd
import re
from rapidfuzz import fuzz, process
from docx import Document
from docx.oxml import OxmlElement
from io import BytesIO
import tempfile
from pathlib import Path
import zipfile
import shutil
from docx.shared import RGBColor
from docxcompose.composer import Composer
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt

# ==========================================
#  1. SESSION & FILE MANAGEMENT
# ==========================================

def get_session_dir():
    """Creates a temporary directory that PERSISTS as long as the session is alive."""
    if "session_work_dir" not in st.session_state:
        temp_dir = tempfile.mkdtemp(prefix="srd_session_")
        st.session_state["session_work_dir"] = temp_dir
    return Path(st.session_state["session_work_dir"])

def split_files_to_disk(files, out_dir, base_name, max_part_mb=80):
    """Splits files into ZIP parts on DISK (low RAM usage)."""
    out_dir = Path(out_dir)
    files = sorted([Path(f) for f in files if Path(f).is_file()], key=lambda p: p.name.lower())
    
    if not files:
        return []

    max_bytes = max_part_mb * 1024 * 1024
    parts = []
    
    part_idx = 1
    current_zip_path = out_dir / f"{base_name}_part{part_idx:02d}.zip"
    
    zf = zipfile.ZipFile(current_zip_path, "w", compression=zipfile.ZIP_DEFLATED)
    current_size = 0
    
    for fp in files:
        fsize = fp.stat().st_size
        if current_size > 0 and (current_size + fsize) > max_bytes:
            zf.close()
            parts.append(current_zip_path)
            part_idx += 1
            current_zip_path = out_dir / f"{base_name}_part{part_idx:02d}.zip"
            zf = zipfile.ZipFile(current_zip_path, "w", compression=zipfile.ZIP_DEFLATED)
            current_size = 0
            
        zf.write(fp, arcname=fp.name)
        current_size += fsize
        
    zf.close()
    parts.append(current_zip_path)
    return parts

# ==========================================
#  2. DOCX & LOGIC HELPER FUNCTIONS
# ==========================================

def recompress_docx_inplace(docx_path: str | Path, remove_thumbnail: bool = True) -> Path:
    docx_path = Path(docx_path)
    if docx_path.suffix.lower() != ".docx": raise ValueError(f"Expected .docx: {docx_path}")
    if not docx_path.exists(): raise FileNotFoundError(docx_path)
    tmp_path = docx_path.with_suffix(".recompressed.tmp")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as zout:
        for item in zin.infolist():
            name = item.filename
            if remove_thumbnail and name.lower() == "docprops/thumbnail.jpeg": continue
            zout.writestr(name, zin.read(name))
    tmp_path.replace(docx_path)
    return docx_path

def force_document_font(doc, font_name="Arial", font_size=12):
    try:
        doc.styles["Normal"].font.name = font_name
        doc.styles["Normal"].font.size = Pt(font_size)
    except: pass
    for paragraph in doc.paragraphs:
        try:
            paragraph.style.font.name = font_name
            paragraph.style.font.size = Pt(font_size)
        except: pass
        for run in paragraph.runs:
            if not (bool(run._r.xpath(".//w:drawing")) or bool(run._r.xpath(".//w:pict"))):
                run.font.size = Pt(font_size)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if not (bool(run._r.xpath(".//w:drawing")) or bool(run._r.xpath(".//w:pict"))):
                            run.font.name = font_name
                            run.font.size = Pt(font_size)

def clean_whitespace(doc):
    for p in list(doc.paragraphs):
        text = "".join(run.text for run in p.runs).strip()
        has_img = bool(p._element.xpath(".//w:drawing")) or bool(p._element.xpath(".//w:pict"))
        has_tbl = bool(p._element.xpath(".//w:tbl"))
        has_br = any(run._r.xpath(".//w:br[@w:type='page']") for run in p.runs)
        if text == "" and not (has_img or has_tbl or has_br):
            p._element.getparent().remove(p._element)

def merge_docx_files(doc_paths, output_path, font_name="Arial", font_size=12):
    if not doc_paths: return
    master = Document(doc_paths[0])
    composer = Composer(master)
    for doc_path in doc_paths[1:]:
        master.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        composer.append(Document(doc_path))
    composer.save(output_path)
    merged = Document(output_path)
    force_document_font(merged, font_name, font_size)
    merged.save(output_path)

def create_reviewer_docx_packets_files(assignments_df, processed_dir) -> list[Path]:
    processed_dir = Path(processed_dir)
    out_files = []
    for reviewer, group in assignments_df.groupby("reviewer_name"):
        abstract_nums = sorted(group["abstract_id"].dropna().astype(int).tolist())
        doc_paths = []
        filename_parts = []
        for num in abstract_nums:
            doc_path = processed_dir / f"srd_abstract_{num}.docx"
            if doc_path.exists():
                doc_paths.append(doc_path)
                filename_parts.append(str(num))
            else:
                filename_parts.append(f"missing{num}")
        if not doc_paths: continue
        nums_str = "-".join(filename_parts)
        out_file = f"{reviewer}_Abstracts_{nums_str}.docx"
        out_path = processed_dir / out_file
        merge_docx_files(doc_paths, output_path=out_path)
        recompress_docx_inplace(out_path, remove_thumbnail=True)
        out_files.append(out_path)
    return out_files

def clean_name(name):
    if name is None: return None
    name = re.sub(r"\^[A-Za-z0-9]+", "", name)
    name = re.sub(r"\^$", "", name)
    name = name.replace("^", "").replace("*", "")
    name = re.sub(r"\d+$", "", name)
    return re.sub(r"\s+", " ", name).strip()

def extract_docx_text_with_superscripts(filepath):
    doc = Document(filepath)
    lines = []
    for p in doc.paragraphs:
        line = ""
        for r in p.runs:
            val = r._element.find(".//w:vertAlign", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
            if val is not None and val.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "superscript":
                line += "^" + r.text
            else:
                line += r.text
        if line.strip(): lines.append(line.strip())
    return lines

MARKER_TIERS = [
    ["choose your research type", "kies uw onderzoekstype", "your research type", "research type:", "onderzoekstype:", "type of research:"],
    ["research classification", "category of research", "clinical research", "fundamental research", "translational research", "basic research", "applied research"],
    ["research type", "type of research", "type onderzoek", "onderzoek type"],
]
FALLBACK_WORDS = ["clinical", "fundamental"]

def _looks_like_option_line(text: str) -> bool:
    t = text.strip()
    if len(t) > 40: return False
    if re.search(r"\b(department|universit|erasmus|mc|rotterdam|affiliation)\b", t, flags=re.I): return False
    return True

def find_first_marker(doc):
    for tier_idx, tier in enumerate(MARKER_TIERS, start=1):
        for i, p in enumerate(doc.paragraphs):
            low = p.text.strip().lower()
            if not low: continue
            for marker in tier:
                if marker in low: return {"paragraph_index": i, "marker": marker, "tier": tier_idx}
    for i, p in enumerate(doc.paragraphs):
        low = p.text.strip().lower()
        if not low or not _looks_like_option_line(p.text): continue
        for w in FALLBACK_WORDS:
            if re.search(rf"\b{re.escape(w)}\b", low): return {"paragraph_index": i, "marker": w, "tier": 4}
    return None

def extract_surname(name):
    if not isinstance(name, str): return None
    clean = re.sub(r"\b(MD|PhD|MSc|BSc|Dr|Prof|Ing|MBA|MPH)\b\.?", "", re.sub(r"[\d\^]+$", "", name.strip()), flags=re.I).strip()
    parts = clean.split()
    if not parts: return None
    if len(parts) == 1: return parts[0].lower()
    last_two = " ".join(parts[-2:]).lower()
    if last_two in {"van", "de", "der", "den", "het", "ter", "ten", "van de", "van der", "van den"}: return last_two
    return parts[-1].lower()

def extract_initials(name):
    if not isinstance(name, str): return None
    name = re.sub(r"\b(MD|PhD|MSc|BSc|Dr|Prof|Ing|MBA|MPH)\b\.?", "", name, flags=re.I).strip()
    dotted = re.findall(r"([A-Za-z])\.", name)
    if dotted: return "".join(d.upper() for d in dotted)
    return "".join(t[0].upper() for t in re.split(r"[ \-]+", name) if t and t[0].isalpha())

def match_author_name(input_name, ref_df, ref_col="name"):
    if pd.isna(input_name) or len(str(input_name).strip()) < 2: return None
    s = str(input_name).strip().replace(",", " ")
    s = re.sub(r"\s+", " ", s).strip()
    candidates = [str(c).strip() for c in ref_df[ref_col].dropna().unique()]
    
    for c in candidates:
        if c.lower() == s.lower(): return c

    hypotheses = [(extract_surname(s), extract_initials(s))]
    tokens = s.split()
    if len(tokens) == 2 and re.fullmatch(r"[A-Za-z]\.?", tokens[1]):
        hypotheses.append((tokens[0].lower(), tokens[1].replace(".", "").upper()))

    best_match = None
    best_score = 0
    
    for in_surname, in_initials in hypotheses:
        if not in_surname: continue
        same_surname = [c for c in candidates if extract_surname(c) == in_surname]
        if len(same_surname) == 1: return same_surname[0]
        
        for c in same_surname:
            score = 0
            if in_initials:
                cand_init = extract_initials(c)
                if cand_init:
                    if cand_init == in_initials: score += 200
                    elif cand_init[0] == in_initials[0]: score += 100
                    score += fuzz.ratio(cand_init, in_initials)
            score += int(0.5 * fuzz.partial_ratio(s.lower(), c.lower()))
            if score > best_score:
                best_score = score
                best_match = c
    
    if best_match and best_score >= 150: return best_match
    
    # Fuzzy Fallback
    match = process.extractOne(s, candidates, scorer=fuzz.partial_ratio)
    if match and match[1] >= 85: return match[0]
    return None

def fuzzy_merge(df1, df2, key1, key2, threshold=90):
    matches = []
    df2_clean = df2.dropna(subset=[key2]).copy()
    df2_clean[key2] = df2_clean[key2].astype(str)
    for _, row in df1.iterrows():
        val = row[key1]
        if pd.isna(val): continue
        match = process.extractOne(str(val), df2_clean[key2].tolist(), scorer=fuzz.partial_ratio)
        if match and match[1] >= threshold:
            matched_row = df2_clean[df2_clean[key2] == match[0]].iloc[0]
            matches.append({**row.to_dict(), **matched_row.to_dict()})
    return pd.DataFrame(matches)

def read_excel_with_auto_header_from_bytes(data: bytes, sheet_name=0):
    """Your original auto-header function, adapted for in-memory bytes."""
    temp = pd.read_excel(BytesIO(data), sheet_name=sheet_name, header=None)
    header_row = temp.notna().any(axis=1).idxmax()
    df = pd.read_excel(BytesIO(data), sheet_name=sheet_name, header=header_row)
    return df

def extract_ids(x):
    nums = re.findall(r"\d+", str(x))
    return [int(n) for n in nums] if nums else None

def department_conflict(author_deps, reviewer_dep, threshold=85):
    if reviewer_dep is None: return False
    rev_d = str(reviewer_dep).lower().strip()
    for auth_d in author_deps:
        if auth_d and fuzz.partial_ratio(str(auth_d).lower().strip(), rev_d) >= threshold: return True
    return False

def assign_reviewers(authors, reviewers, max_reviews=8, reviewers_per_abs=3, conflict_threshold=85):
    assignments = []
    for _, row in authors.iterrows():
        author_name = row["name"]
        abstract_id = row["abstract_id"]
        author_deps = list(row["departments"])
        
        mask = (reviewers["assigned_count"] < max_reviews) & (~reviewers["reviewer_department"].apply(lambda d: department_conflict(author_deps, d, threshold=conflict_threshold)))
        eligible = reviewers[mask]
        
        if len(eligible) < reviewers_per_abs:
            raise ValueError(f"Not enough reviewers for {author_name} (ID {abstract_id}). Needed {reviewers_per_abs}, found {len(eligible)}.")
        
        chosen = eligible.sort_values(by="assigned_count").head(reviewers_per_abs)
        for i, reviewer in chosen.iterrows():
            reviewers.loc[i, "assigned_count"] += 1
            assignments.append({
                "author_name": author_name, 
                "abstract_id": abstract_id, 
                "reviewer_name": reviewer["reviewer_name"], 
                "reviewer_department": reviewer["reviewer_department"], 
                "reviewer_load_after": reviewers.loc[i, "assigned_count"], 
                "author_departments": ", ".join(author_deps)
            })
    return pd.DataFrame(assignments), reviewers

def prepare_ref_and_authors(ref_file, trans_file):
    trans_df = pd.read_excel(trans_file).drop_duplicates()
    ref_df = pd.read_excel(ref_file)
    ref_df.columns = ref_df.columns.str.lower().str.replace(" ", "_")
    
    for col in ["name", "department_", "abstract_nr._"]:
        if col in ref_df.columns: ref_df[col] = ref_df[col].astype(str).str.strip()
    
    ref_df["department_"] = ref_df["department_"].replace("nan", "").fillna("")
    ref_df["department"] = ref_df["department_"].str.split(r"[;,+/]+", regex=True)
    ref_expanded = ref_df.explode("department")
    ref_expanded["department"] = ref_expanded["department"].str.strip()
    ref_expanded = ref_expanded[ref_expanded["department"] != ""]
    
    ref_merged = fuzzy_merge(ref_expanded, trans_df, "department", "department")
    ref_merged["abstract_id_list"] = ref_merged["abstract_nr._"].apply(extract_ids)
    ref_merged = ref_merged.explode("abstract_id_list").rename(columns={"abstract_id_list": "abstract_id"})
    
    authors = ref_merged.groupby(["name", "abstract_id"])["English"].apply(lambda x: list(set(x.dropna()))).reset_index().rename(columns={"English": "departments"})
    return ref_df, authors

def process_doc(filepath, ref_df, output_folder, remaining_ids):
    txt = extract_docx_text_with_superscripts(filepath)
    def find_line(phrase):
        for i, line in enumerate(txt):
            if phrase.lower() in line.lower(): return i
        return None
    
    author_line = find_line("author")
    aff_line = find_line("affiliations")

    if aff_line is None or author_line is None:
        doc = Document(filepath)
        out = Path(output_folder) / (Path(filepath).stem + "_unchanged.docx")
        doc.save(out)
        recompress_docx_inplace(out)
        return {"file": Path(filepath).name, "matched": False}

    def extract_first_author(idx):
        line = txt[idx].replace("Authors:", "").replace("Author:", "").strip()
        if line and not line.lower().startswith("affiliation"):
            return re.sub(r"\d+$", "", re.split(r"[;,&]", line)[0].strip()).strip()
        for i in range(idx + 1, len(txt)):
            cand = txt[i].strip()
            if not cand: continue
            if "affiliation" in cand.lower(): break
            return re.sub(r"\d+$", "", cand.split(",")[0].strip()).strip()
        return ""

    name = clean_name(extract_first_author(author_line))
    matched_name = match_author_name(name, ref_df)
    
    abstract_nr = None
    if matched_name:
        key = matched_name.strip()
        if key in remaining_ids and remaining_ids[key]:
            abstract_nr = remaining_ids[key].pop(0)

    doc = Document(filepath)
    res = find_first_marker(doc)
    
    if not res:
        out = Path(output_folder) / (Path(filepath).stem + "_unchanged.docx")
        doc.save(out)
        return {"file": Path(filepath).name, "name": name, "matched": abstract_nr is not None}

    # Cut top
    for _ in range(res["paragraph_index"]):
        p = doc.paragraphs[1]
        p._element.getparent().remove(p._element)
    
    clean_whitespace(doc)

    if abstract_nr is not None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Abstract number: {abstract_nr}")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)
        doc._element.body.remove(p._p)
        doc.paragraphs[0]._p.addnext(p._p)

    for _ in range(5): doc._element.body.insert(1, OxmlElement("w:p"))
    force_document_font(doc)

    fname = f"srd_abstract_{abstract_nr}.docx" if abstract_nr else Path(filepath).stem + "_no_number.docx"
    out = Path(output_folder) / fname
    doc.save(out)
    recompress_docx_inplace(out)
    
    return {"file": Path(filepath).name, "name": name, "matched_name": matched_name, "abstract_nr": abstract_nr, "matched": abstract_nr is not None}

def build_ids(ref_df):
    ids = {}
    for _, row in ref_df.iterrows():
        name = str(row.get("name", "")).strip()
        nums = extract_ids(row.get("abstract_nr._", ""))
        if nums:
            if name not in ids: ids[name] = []
            for n in nums: 
                if n not in ids[name]: ids[name].append(n)
            ids[name].sort()
    return ids

# ==========================================
#  3. MAIN PIPELINE
# ==========================================

def run_pipeline(ref_file, trans_file, reviewer_file, docx_files, max_part_mb=80):
    work_dir = get_session_dir()
    input_dir, output_dir, final_dir = work_dir/"input", work_dir/"output", work_dir/"final_zips"
    
    for d in [input_dir, output_dir, final_dir]:
        if d.exists(): shutil.rmtree(d)
        d.mkdir(parents=True)

    # Save Uploads
    for uf in docx_files: (input_dir / uf.name).write_bytes(uf.getvalue())
        
    # Process Logic
    ref_df, authors = prepare_ref_and_authors(ref_file, trans_file)
    remaining_ids = build_ids(ref_df)
    
    reviewer_df = read_excel_auto_header(reviewer_file.getvalue())
    reviewer_df.columns = reviewer_df.columns.str.lower().str.strip()
    reviewer_df = reviewer_df.rename(columns={"reviewer signup": "reviewer_name", "department": "reviewer_department"})
    reviewer_df["assigned_count"] = 0
    reviewer_df = reviewer_df.dropna(subset=["reviewer_name"])
    
    assignments_df, _ = assign_reviewers(authors, reviewer_df)

    # Process Docs
    results = []
    for fpath in input_dir.glob("*.docx"):
        results.append(process_doc(str(fpath), ref_df, str(output_dir), remaining_ids))
    
    # Create ZIPs
    abs_parts = split_files_to_disk(list(output_dir.glob("*.docx")), final_dir, "Abstracts", max_part_mb)
    rev_docs = create_reviewer_docx_packets_files(assignments_df, output_dir)
    rev_parts = split_files_to_disk(rev_docs, final_dir, "ReviewerPackets", max_part_mb)

    assign_path = final_dir / "assignments.xlsx"
    assignments_df.to_excel(assign_path, index=False)

    return assignments_df, pd.DataFrame(results), assign_path, abs_parts, rev_parts

# ==========================================
#  4. STREAMLIT UI
# ==========================================

st.title("SRD Abstracts Processor")
st.caption("Optimized for reliability and large file handling.")

with st.sidebar:
    st.header("1. Upload input files")
    ref_file = st.file_uploader("Registrations overview (Excel)", type=["xlsx", "xls"], key="ref")
    trans_file = st.file_uploader("Department translations (Excel)", type=["xlsx", "xls"], key="trans")
    reviewer_file = st.file_uploader("Reviewers (Excel)", type=["xlsx", "xls"], key="review")
    docx_files = st.file_uploader("Abstract DOCX files (multiple)", type=["docx"], accept_multiple_files=True)

    st.header("2. Run")
    run_btn = st.button("Run processing", type="primary")

# Initialize Session State
if "processed" not in st.session_state:
    st.session_state['processed'] = False
    st.session_state['assignments_df'] = None
    st.session_state['results_df'] = None
    st.session_state['assign_path'] = None
    st.session_state['abs_parts'] = []
    st.session_state['rev_parts'] = []

if run_btn:
    if not (ref_file and trans_file and reviewer_file and docx_files):
        st.error("Please upload all required files.")
    else:
        with st.spinner("Processing... This may take a minute."):
            try:
                assignments_df, results_df, assign_path, abs_parts, rev_parts = run_pipeline(
                    ref_file, trans_file, reviewer_file, docx_files, max_part_mb=80
                )
                
                st.session_state['processed'] = True
                st.session_state['assignments_df'] = assignments_df
                st.session_state['results_df'] = results_df
                st.session_state['assign_path'] = str(assign_path)
                st.session_state['abs_parts'] = [str(p) for p in abs_parts]
                st.session_state['rev_parts'] = [str(p) for p in rev_parts]
                
                st.success("Processing Complete!")
            except Exception as e:
                st.error(f"Error: {e}")

# --- DISPLAY OUTPUTS ---
if st.session_state.get('processed'):
    st.divider()
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Assignments")
        if st.session_state['assignments_df'] is not None:
            st.dataframe(st.session_state['assignments_df'], height=300, use_container_width=True)
        
        if st.session_state.get('assign_path'):
            path = Path(st.session_state['assign_path'])
            if path.exists():
                st.download_button(
                    label="📊 Download Assignments Excel", 
                    data=path.read_bytes(), 
                    file_name="assignments.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )

    with col2:
        st.subheader("Processing Log")
        if st.session_state['results_df'] is not None:
            st.dataframe(st.session_state['results_df'], height=300, use_container_width=True)

    st.divider()

    # --- NEW: FILE EXPLANATIONS ---
    st.header("📄 File Guide")
    st.markdown("""
    * **Assignments Excel:** A spreadsheet listing which reviewer is assigned to which abstract.
    * **Abstracts_partXX.zip:** Contains the individual, processed DOCX files for all authors. Use these if you need to access specific abstracts one-by-one.
    * **ReviewerPackets_partXX.zip:** Contains merged DOCX files for each reviewer. Each file (e.g., `ReviewerName_Abstracts_01-05.docx`) contains all the abstracts assigned to that specific reviewer combined into one document.
    """)

    # --- NEW: HOW TO DOWNLOAD ---
    st.header("📥 Download ZIP Files")
    st.info("""
    **How to download:**
    1. Click the dropdown menu below.
    2. Select the file you want (e.g., `ReviewerPackets_part01.zip`).
    3. Wait a moment for the specific "Download" button to appear.
    4. Click the button to save the file.
    
    *Note: Downloading files one by one prevents the application from crashing due to memory limits.*
    """)

    all_parts = st.session_state['abs_parts'] + st.session_state['rev_parts']
    parts_map = {Path(p).name: p for p in all_parts}
    options = list(parts_map.keys())
    
    selected_file_name = st.selectbox("Select file to download:", options=options)
    
    if selected_file_name:
        full_path_str = parts_map[selected_file_name]
        full_path = Path(full_path_str)
        
        if full_path.exists():
            file_size_mb = round(full_path.stat().st_size / (1024 * 1024), 2)
            with open(full_path, "rb") as f:
                st.download_button(
                    label=f"⬇️ Download {selected_file_name} ({file_size_mb} MB)",
                    data=f,
                    file_name=selected_file_name,
                    mime="application/zip",
                    key="dynamic_dl_btn"
                )
        else:
            st.error("File not found on server. The session may have expired.")
